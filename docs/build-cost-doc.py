"""Build docs/inventory-cost-accounting.md -> a professional Word (.docx).
Handles headings, paragraphs, **bold**/`code`, tables, ```code fences```,
> blockquotes, - bullets, 1. ordered lists, --- rules. python-docx only.
"""
import os, re, datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DIR = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(DIR, "inventory-cost-accounting.md")
OUT = os.path.join(DIR, "inventory-cost-accounting.docx")

ACCENT = RGBColor(0xF9, 0x73, 0x16)
MUTED  = RGBColor(0x64, 0x74, 0x8B)
INK    = RGBColor(0x0F, 0x17, 0x2A)
BROWN  = RGBColor(0x7C, 0x2D, 0x12)
CODEBG = "F1F5F9"; QUOTEBG = "FFF7ED"; HDRFILL = "F97316"
USABLE = 6.7  # inches (A4 minus margins)

INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")

def add_runs(p, text, base_color=None, mono_size=9.5):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
            if base_color: r.font.color.rgb = base_color
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(mono_size)
            r.font.color.rgb = RGBColor(0x0B, 0x52, 0x94)
        else:
            r = p.add_run(tok)
            if base_color: r.font.color.rgb = base_color

def shade_para(p, fill):
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    p._p.get_or_add_pPr().append(shd)

def shade_cell(cell, fill):
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)

def set_col_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)

def is_cont(s):
    """A wrapped continuation line of the current list item (not a new block)."""
    return bool(s.strip()) and not re.match(r"^(#{1,6}\s|>\s?|\s*[-*]\s|\s*\d+\.\s|---+\s*$|\s*\|)", s) and not s.lstrip().startswith("```")

doc = Document()
style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6); style.paragraph_format.line_spacing = 1.15
for sec in doc.sections:
    sec.left_margin = sec.right_margin = Inches(0.7)
    sec.top_margin = sec.bottom_margin = Inches(0.8)

# footer
footer = doc.sections[0].footer
fp = footer.paragraphs[0]; fp.text = "Alfaytri — Inventory Cost Accounting"
fp.runs[0].font.size = Pt(8); fp.runs[0].font.color.rgb = MUTED

lines = open(SRC, encoding="utf-8").read().split("\n")
i = 0
first_h1_done = False

def is_table_sep(s):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", s)) and "-" in s

def split_row(s):
    s = s.strip()
    if s.startswith("|"): s = s[1:]
    if s.endswith("|"): s = s[:-1]
    return [c.strip() for c in s.split("|")]

while i < len(lines):
    line = lines[i]
    if not line.strip():
        i += 1; continue

    # code fence
    if line.lstrip().startswith("```"):
        i += 1; buf = []
        while i < len(lines) and not lines[i].lstrip().startswith("```"):
            buf.append(lines[i]); i += 1
        i += 1  # closing fence
        t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = t.cell(0, 0); shade_cell(cell, CODEBG)
        cell.paragraphs[0].text = ""
        for k, cl in enumerate(buf):
            p = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
            r = p.add_run(cl if cl else " "); r.font.name = "Consolas"; r.font.size = Pt(9)
        set_col_widths(t, [USABLE])
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        continue

    # table
    if line.strip().startswith("|") and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
        header = split_row(line); i += 2
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            rows.append(split_row(lines[i])); i += 1
        ncol = len(header)
        t = doc.add_table(rows=1, cols=ncol); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, htext in enumerate(header):
            c = t.rows[0].cells[j]; shade_cell(c, HDRFILL)
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
            for tok in INLINE.split(htext):
                if not tok: continue
                txt = tok[2:-2] if tok.startswith("**") else (tok[1:-1] if tok.startswith("`") else tok)
                r = p.add_run(txt); r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r.font.size = Pt(9.5)
        for rowvals in rows:
            cells = t.add_row().cells
            for j in range(ncol):
                val = rowvals[j] if j < len(rowvals) else ""
                p = cells[j].paragraphs[0]; p.paragraph_format.space_after = Pt(2)
                add_runs(p, val, mono_size=9)
                for r in p.runs: r.font.size = r.font.size or Pt(9.5)
        widths = [USABLE / ncol] * ncol
        # give a wider first column for 2-3 col descriptive tables
        if ncol <= 3:
            widths = [USABLE * 0.32] + [ (USABLE*0.68)/(ncol-1) ]*(ncol-1)
        set_col_widths(t, widths)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        continue

    # heading
    mh = re.match(r"^(#{1,6})\s+(.*)$", line)
    if mh:
        lvl = len(mh.group(1)); txt = mh.group(2)
        if lvl == 1 and not first_h1_done:
            first_h1_done = True
            tp = doc.add_paragraph(); tp.paragraph_format.space_after = Pt(2)
            r = tp.add_run(txt); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = ACCENT
            sub = doc.add_paragraph();
            rs = sub.add_run("Alfaytri — Technical Reference  ·  Generated " + datetime.date.today().isoformat())
            rs.font.size = Pt(9.5); rs.font.color.rgb = MUTED
            bd = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"18"); bot.set(qn("w:space"),"6"); bot.set(qn("w:color"),"F97316")
            bd.append(bot); sub._p.get_or_add_pPr().append(bd)
            i += 1; continue
        h = doc.add_heading(level=min(lvl, 3))
        add_runs(h, re.sub(r"^\d+\.\s*", "", txt) if False else txt)
        if lvl == 1:
            for r in h.runs: r.font.color.rgb = ACCENT
        i += 1; continue

    # hr
    if re.match(r"^---+\s*$", line):
        i += 1; continue

    # blockquote
    if re.match(r"^>\s?", line):
        buf = []
        while i < len(lines) and re.match(r"^>\s?", lines[i]):
            buf.append(re.sub(r"^>\s?", "", lines[i])); i += 1
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
        shade_para(p, QUOTEBG)
        add_runs(p, " ".join(x for x in buf if x.strip()), base_color=BROWN)
        continue

    # ordered list — explicit numbers (reliable; Word's List Number can continue a prior list)
    if re.match(r"^\s*\d+\.\s+", line):
        n = 0
        while i < len(lines):
            if re.match(r"^\s*\d+\.\s+", lines[i]):
                n += 1
                txt = re.sub(r"^\s*\d+\.\s+", "", lines[i]); i += 1
                while i < len(lines) and is_cont(lines[i]):
                    txt += " " + lines[i].strip(); i += 1
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.35)
                p.paragraph_format.first_line_indent = Inches(-0.22)
                p.paragraph_format.space_after = Pt(3)
                rn = p.add_run(f"{n}. "); rn.bold = True
                add_runs(p, txt)
            else:
                break
        continue

    # bullet list
    if re.match(r"^\s*[-*]\s+", line):
        while i < len(lines):
            if re.match(r"^\s*[-*]\s+", lines[i]):
                indent = len(re.match(r"^(\s*)", lines[i]).group(1))
                txt = re.sub(r"^\s*[-*]\s+", "", lines[i]); i += 1
                while i < len(lines) and is_cont(lines[i]):
                    txt += " " + lines[i].strip(); i += 1
                p = doc.add_paragraph(style="List Bullet" if indent < 2 else "List Bullet 2")
                add_runs(p, txt)
            else:
                break
        continue

    # paragraph
    buf = [line]; i += 1
    while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,6}\s|>\s?|\s*[-*]\s|\s*\d+\.\s|---+\s*$|\s*\|)", lines[i]) and not lines[i].lstrip().startswith("```"):
        buf.append(lines[i]); i += 1
    p = doc.add_paragraph(); add_runs(p, " ".join(buf))

doc.save(OUT)
print(f"DOCX written: {OUT} ({round(os.path.getsize(OUT)/1024)} KB)")
