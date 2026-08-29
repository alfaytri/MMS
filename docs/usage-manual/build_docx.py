"""Builds the Alfaytri Staff User Guide as a Word (.docx) from the ordered
Markdown sections in this folder, embedding the screenshots in assets/.
Single source of truth = the same .md files the PDF builder uses.
Run from repo root:  python docs/usage-manual/build_docx.py
"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DIR = os.path.dirname(os.path.abspath(__file__))
SECTIONS = ["00-cover", "01-getting-started", "02-master-data", "03-purchase", "04-sales", "05-operations", "06-reports", "07-common-tasks"]

ACCENT = RGBColor(0xF9, 0x73, 0x16)
MUTED = RGBColor(0x64, 0x74, 0x8B)
INK = RGBColor(0x0F, 0x17, 0x2A)
USABLE_WIDTH_IN = 6.3  # A4 minus ~1in margins each side

INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_runs(paragraph, text):
    """Add text to a paragraph, honouring **bold** and `code` inline markers."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Consolas"; r.font.size = Pt(10)
        else:
            paragraph.add_run(tok)


def shade(paragraph, hexfill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexfill)
    pPr.append(shd)


doc = Document()
# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"; style.font.size = Pt(11)
for sec in doc.sections:
    sec.left_margin = sec.right_margin = Inches(0.85)
    sec.top_margin = sec.bottom_margin = Inches(0.8)

img_count = 0
for idx, name in enumerate(SECTIONS):
    md = open(os.path.join(DIR, name + ".md"), encoding="utf-8").read()
    lines = md.split("\n")
    i = 0
    if idx > 0:
        doc.add_page_break()
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1; continue
        m_img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
        m_h = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m_img:
            alt, src = m_img.group(1), m_img.group(2)
            path = os.path.join(DIR, src)
            if os.path.exists(path):
                doc.add_picture(path, width=Inches(USABLE_WIDTH_IN))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_count += 1
                if alt:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = cap.add_run(alt); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MUTED
            i += 1; continue
        if re.match(r"^---\s*$", line):
            i += 1; continue  # decorative rule -> skip in Word
        if m_h:
            level = len(m_h.group(1)); txt = m_h.group(2)
            h = doc.add_heading(level=min(level, 3))
            add_runs(h, txt)
            if level == 1:
                for r in h.runs:
                    r.font.color.rgb = ACCENT
            i += 1; continue
        if re.match(r"^>\s?", line):
            buf = []
            while i < len(lines) and re.match(r"^>\s?", lines[i]):
                buf.append(re.sub(r"^>\s?", "", lines[i])); i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
            shade(p, "FFF7ED")
            add_runs(p, " ".join(buf))
            for r in p.runs:
                r.font.color.rgb = RGBColor(0x7C, 0x2D, 0x12)
            continue
        if re.match(r"^\s*-\s+", line):
            while i < len(lines) and re.match(r"^\s*-\s+", lines[i]):
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, re.sub(r"^\s*-\s+", "", lines[i])); i += 1
            continue
        # paragraph
        buf = [line]; i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,4}\s|!\[|>\s?|\s*-\s|---\s*$)", lines[i]):
            buf.append(lines[i]); i += 1
        p = doc.add_paragraph()
        add_runs(p, " ".join(buf))

out = os.path.join(DIR, "Alfaytri-User-Guide.docx")
doc.save(out)
print(f"DOCX written: {out} ({round(os.path.getsize(out)/1024)} KB, {img_count} images)")
